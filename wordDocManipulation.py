import boto3
import searchText
import docx
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.dml import MSO_THEME_COLOR

import os.path
import datetime


#-----------------------------------
def combineWordParaToString(word_file):
    doc = docx.Document(word_file)
    text = '\n'.join([p.text for p in doc.paragraphs])
    return text

#---------------------------------

'''Recieves wd and returns set of key phrases - returning a set avoids duplication of key phrases'''
def getKeyPhrasesFromDoc(word_file):
    text = combineWordParaToString(word_file)
    key_phrases = searchText.getKeyPhrases(text)
    return key_phrases
#----------------------------------------------------------

'''Get word count from a word doc'''
def getDocWordCount(word_file):
    document = docx.Document(word_file)
    wordCount = 0
    for para in document.paragraphs:
        wordCount += len(para.text.split())
    return wordCount
# ---------------------------------------------------------
'''Get date created from a word doc'''
def getDocDateCreated(word_file):
    document = docx.Document(word_file)
    date = document.core_properties.created
    return date
# ---------------------------------------------------------
''''Get date modified from a word doc'''
def getDocDateModified(word_file):
    document = docx.Document(word_file)
    date = document.core_properties.modified
    return date

# ---------------------------------------------
def createFormattedBrick(RoughDraftDoc):
    return 1
#-------------------------------------------------------

def highlightKeyPoints(sourceDocument):
    return 1
#-------------------------------------------------------
def highlightRedundancies(sourceDocument):
    return 1
#-------------------------------------------------------
def highlightContradictions(sourceDocument):
    return 1
#-------------------------------------------------------
def addPicture(picture):
    return 1
#-------------------------------------------------------

'''Extract headings from a word doc'''
def extractHeadings(word_file):
    doc = docx.Document(word_file)
    paragraphs = doc.paragraphs
    headings = []
    for paragraph in paragraphs:
        if paragraph.style.name.startswith('Heading'):
            headings.append(paragraph.text)
    return headings
#--------------------------------

'''Extract bulleted list items from a word doc'''
def extractBullets(word_file):
    doc = docx.Document(word_file)
    bulletPoints = []
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('List Bullet'):
            bulletPoints.append(paragraph.text)
    return bulletPoints
# -----------------------------

'''Create a template formated word doc (initiate format to Brick)'''
def createBrick():
    Brick = docx.Document()
    Brick.add_heading('CRISPR Gene Editing', level=1)
    Brick.add_heading('Objectives', level=1)
    Brick.add_paragraph(style='List Bullet')
    for item in['basic DNA structure', 'a viral defense mechanism', 'clustered regularly interspaced short palindromic repeates', 'Cas9', 'CRISPR-Cas9 for DNA editing']:
        Brick.add_paragraph(item, style='List Bullet')
    
    Brick.add_heading('A review of DNA structure', level=2)
    Brick.add_paragraph('This is the first section')
    
    Brick.add_heading('Do Bacteria have immune systems?', level=2)
    Brick.add_paragraph('This is the second section')
    
    Brick.add_heading('What are CRISPR?', level=2)
    Brick.add_heading('What is the Cas9 protein', level=2)
    Brick.add_heading('CRISPR-Cas9: Uses and Applications', level=2)
    
    Brick.add_heading('Summary', level=1)
    Brick.add_heading('Questions', level=1)
    Brick.save('Brick0.docx')  
#--------------------------------
