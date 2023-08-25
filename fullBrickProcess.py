import io
import os

import boto3

import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from pdf2image import convert_from_path
from PyPDF2 import PdfReader


#local imports
import WordDocx_Summerizer as summerizer

#---------------------------------------------------------------------------------------
# MISC. FUNCTIONS COPIED OVER FROM ELSEWHERE TO REDUCE NUMBER OF IMPORTS AND TO CONDENSE
# --------------------------------------------------------------------------------------
'''Given a string of text, returns a set of key phrases, using comprehend'''
def getKeyPhrases(text):
    comprehend = boto3.client('comprehend')
    keyPhrases = comprehend.detect_key_phrases(Text=text, LanguageCode='en')
    keyPhraseSet = set()
    for phrase in keyPhrases['KeyPhrases']:
        keyPhraseSet.add(phrase['Text'])
    return keyPhraseSet

'''Combines all the text in the word document into one string of words'''
def combineWordParaToString(word_file):
    doc = docx.Document(word_file)
    text = '\n'.join([p.text for p in doc.paragraphs])
    return text

'''Recieves wd and returns set of key phrases - returning a set avoids duplication of key phrases'''
def getKeyPhrasesFromDoc(word_file):
    text = combineWordParaToString(word_file)
    key_phrases = getKeyPhrases(text)
    return key_phrases
# ---------------------------------
''' Takes all the text from the pdf file and dumps it into a word file'''
def dump_pdfText2Word(pdf_file, word_file):
    
    # Create a new word document
    brickDraft = docx.Document()
    
    # Read the pdf file
    pdf_reader = PdfReader(pdf_file)
    
    for page in pdf_reader.pages:
        text = page.extract_text()
        brickDraft.add_paragraph(text)

    brickDraft.save(word_file)

# ---------------------------------
'''Separates references from the rest of the source material'''
def separateReferences(wordSource, referencesName, draftMaterialName, keyword):
    
    # Read the word file
    mainDoc = docx.Document(wordSource)
    title = str(mainDoc.core_properties.title)
    
    references = docx.Document()
    brickDraftMaterial = docx.Document()
    
    
    for i, para in enumerate(mainDoc.paragraphs):
        if keyword in para.text:
            for j in range(i, len(mainDoc.paragraphs)):
                references.add_paragraph(mainDoc.paragraphs[j].text)
            for k in range(0, i):
                brickDraftMaterial.add_paragraph(mainDoc.paragraphs[k].text)
    references.save(referencesName)
    brickDraftMaterial.save(draftMaterialName)

'''Extract text between two key words in a word doc'''
def extractTextbetweenHeadings(docName, heading1, heading2):
    document = Document(docName)
    text = ""
    found = False
    for paragraph in document.paragraphs:
        if paragraph.text == heading1:
            found = True
        elif paragraph.text == heading2:
            break
        elif found:
            text += paragraph.text    
    return text

def addTextAfter(document, heading, paragraph_text):
    document = docx.Document(document)
    for paragraph in document.paragraphs:
        if paragraph.text == heading:
            new_paragraph = document.add_paragraph(paragraph_text)
            break
    document.save(document)
# ---------------------------------
'''Gets all images within a pdf file'''     
def get_pdf_images(pdf_file):
    reader = PdfReader(pdf_file)
    page = reader.pages[0]
    count = 0

    for image_file_object in page.images:
        with open(str(count) + image_file_object.name, "wb") as fp:
            fp.write(image_file_object.data)
            count += 1
# ---------------------------------------
'''Add a centered heading'''
# -----------------------------------
'''Summarize the source for the draft material'''
def appendSummary(wordDocument):
    # Read the word file
    mainDoc = docx.Document(wordDocument)
    summary = summerizer.summarize_docx_file(wordDocument)
    
    heading = mainDoc.add_heading('Summary', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mainDoc.add_paragraph(summary)
    mainDoc.save(wordDocument)

# -----------------------------------------
def appendKeyWords(wordDocument):
    mainDoc = docx.Document(wordDocument)

    heading = mainDoc.add_heading('Key Words', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    keyPhrases = getKeyPhrasesFromDoc(wordDocument)
    para = mainDoc.add_paragraph()
    phraseString = ', '.join(keyPhrases)
    para.add_run(phraseString)

    mainDoc.save(wordDocument)
# -----------------------------------
'''Create a template formated word doc (initiate format to Brick)'''
def createBrick(brickTitle):
    Brick = docx.Document()
    Brick.add_heading('Title', level=1)
    Brick.add_heading('Objectives', level=1)
    Brick.add_paragraph(style='List Bullet')
    for item in['objective1', 'objective2', 'objective3', 'objective4', 'objective5']:
        Brick.add_paragraph(item, style='List Bullet')
    
    Brick.add_heading('Text Blob', level=1)
    
    for item in['objective1', 'objective2', 'objective3', 'objective4', 'objective5']:
        Brick.add_heading(item, level=2)
        Brick.add_paragraph('Here is objective relavent text')
    
    Brick.add_heading('Summary', level=1)
    
    Brick.add_heading('Questions', level=1)
    
    Brick.add_heading('Sources', level=1)

    Brick.save(brickTitle)  
# -----------------------------------
'''Extract and combine summaries into a brick draft'''
def extractAndFormatBrick(document, brickDraft):
    
    brickDraft = docx.Document(brickDraft)
    document = docx.Document(document)
    text = extractTextbetweenHeadings(document, "Summary", "Key Words")
    addTextAfter(brickDraft, "Text Blob", text)
        
    brickDraft.save(brickDraft) 


    
# get_pdf_images("PDF_Sources/THC.pdf")

'''
dump_pdfText2Word("PDF_Sources/THC.pdf", "Word_Sources/THC.docx")
dump_pdfText2Word("PDF_Sources/IMH.pdf", "Word_Sources/IMH.docx")
dump_pdfText2Word("PDF_Sources/RMC.pdf", "Word_Sources/RMC.docx")

separateReferences("Word_Sources/THC.docx", "Word_Sources/THCR.docx", "Word_Sources/THCDM.docx", "References")
separateReferences("Word_Sources/IMH.docx", "Word_Sources/IMHR.docx", "Word_Sources/IMHDM.docx", "References")
separateReferences("Word_Sources/RMC.docx", "Word_Sources/RMCR.docx", "Word_Sources/RMCDM.docx", "References")

appendSummary("Word_Sources/THCDM.docx")
appendSummary("Word_Sources/IMHDM.docx")
appendSummary("Word_Sources/RMCDM.docx")

appendKeyWords("Word_Sources/THCDM.docx")
appendKeyWords("Word_Sources/IMHDM.docx")
appendKeyWords("Word_Sources/RMCDM.docx")
'''
separateReferences("PDF_Sources/THC.docx", "PDF_Sources/THCR.docx", "PDF_Sources/THCDM.docx", "References")
appendSummary("PDF_Sources/THCDM.docx")
appendKeyWords("PDF_Sources/THCDM.docx")

separateReferences("PDF_Sources/IMH.docx", "PDF_Sources/IMHR.docx", "PDF_Sources/IMHDM.docx", "References")
appendSummary("PDF_Sources/IMHDM.docx")
appendKeyWords("PDF_Sources/IMHDM.docx")

createBrick('Bricks/Brick0.docx')

text = extractTextbetweenHeadings('PDF_Sources/THCDM.docx', "Summary", "Key Words")
brickDraft = docx.Document('Bricks/Brick0.docx')
brickDraft.add_paragraph(text)
brickDraft.save('Bricks/Brick0.docx')