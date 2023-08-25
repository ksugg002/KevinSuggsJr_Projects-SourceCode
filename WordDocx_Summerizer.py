import spacy
from spacy.lang.en.stop_words import STOP_WORDS
from string import punctuation
from collections import Counter
from heapq import nlargest
import docx
import os


def get_text_from_docx(file_path):
    # open the Word document
    doc = docx.Document(file_path)

    # create an empty string to store the text
    text = ''

    # loop through each paragraph in the document and add its text to the string
    for paragraph in doc.paragraphs:
        text += paragraph.text

    return text


def preprocess_text(text):
    nlp = spacy.load("en_core_web_sm")
    doc = nlp(text)

    # Calculate the frequency of each token using the “Counter” function. Top 5 most frequent words.
    keyword = []
    stopwords = list(STOP_WORDS)
    pos_tag = ['PROPN', 'ADJ', 'NOUN', 'VERB']
    for token in doc:
        if (token.text in stopwords or token.text in punctuation):
            continue
        if (token.pos_ in pos_tag):
            keyword.append(token.text)
    freq_word = Counter(keyword)

    # result is stored as a key-value pair in sent_strength where keys are the sentences in the string doc and the values are the weight of each sentence
    max_freq = Counter(keyword).most_common(1)[0][1]
    for word in freq_word.keys():
        freq_word[word] = (freq_word[word]/max_freq)

    return doc, freq_word


def calculate_sentence_strength(doc, freq_word):
    sent_strength = {}
    for sent in doc.sents:
        for word in sent:
            if word.text in freq_word.keys():
                if sent in sent_strength.keys():
                    sent_strength[sent] += freq_word[word.text]
                else:
                    sent_strength[sent] = freq_word[word.text]
    return sent_strength


def summarize_doc(doc, sent_strength, num_sentences):
    # nlargest function is used to summarize the string
    summarized_sentences = nlargest(
        num_sentences, sent_strength, key=sent_strength.get)

    # finalized output
    final_sentences = [w.text for w in summarized_sentences]
    summary = ' '.join(final_sentences)
    return summary


def summarize_docx_file(file_path, num_sentences=3):
    text = get_text_from_docx(file_path)
    doc, freq_word = preprocess_text(text)
    sent_strength = calculate_sentence_strength(doc, freq_word)
    summary = summarize_doc(doc, sent_strength, num_sentences)
    return summary


if __name__ == "__main__":
    doc = docx.Document(os.path.abspath('/Users/kevinsuggs/Desktop/cs410w-411/s23-Black-1-9/BM_summerizer_code/MuscularSystems-2.docx'))
    # loop through each paragraph in the document and add its text to the string
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text

    doc = text


    summary = summarize_docx_file(doc)
    print(summary)

